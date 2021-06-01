import csv
import os
from io import BytesIO

import xlsxwriter
from django.conf import settings
from django.contrib.humanize.templatetags.humanize import naturaltime
from django.db.models import Q, Sum
from django.http import HttpResponse
from django.template.defaultfilters import pluralize, slugify, date
from django.utils import timezone
from django.utils.translation import gettext as _
from docx import Document
from html2text import html2text
from openpyxl import load_workbook

from lib.functions.custom_functions import listrify
from lib.templatetags.custom_filters import nz, currency
from lib.templatetags.verbose_names import get_verbose_label, get_field_value
from shared_models.models import Region, FiscalYear, Section
from . import models, utils
from .models import ProjectYear
from .utils import in_projects_admin_group

from .api.views import get_project_year_queryset
from shared_models import models as shared_models


def generate_acrdp_application(project):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.docx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "acrdp_template.docx")

    with open(template_file_path, 'rb') as f:
        source_stream = BytesIO(f.read())
    document = Document(source_stream)
    source_stream.close()

    lead = None
    contact_info = None
    if project.lead_staff.exists():
        lead = project.lead_staff.first().user
        contact_info = _("{full_address}\n\n{email}\n\n{phone}").format(
            full_address=project.organization.full_address if project.organization else "MISSING!",
            email=lead.email,
            phone=nz(lead.profile.phone, "MISSING!")
        )

    priorities = str()
    for year in project.years.all():
        priorities += f'{year.fiscal_year}:\n\n{year.priorities}\n\n'

    deliverables = str()
    for year in project.years.all():
        if year.activities.filter(type=2).exists():
            deliverables += f'{year.fiscal_year}:\n\n'
            i = 1
            for d in year.activities.filter(type=2):
                deliverables += f'{i}) {d.name.upper()} - {d.description}\n\n'
                i += 1

    field_dict = dict(
        TAG_TITLE=project.title,
        TAG_ORG_NAME=project.organization.tname if project.organization else "MISSING!",
        TAG_ADDRESS=project.organization.address if project.organization else "MISSING!",
        TAG_CITY=project.organization.city if project.organization else "MISSING!",
        TAG_PROV=str(project.organization.location.tname) if project.organization else "MISSING!",
        TAG_POSTAL_CODE=project.organization.postal_code if project.organization else "MISSING!",
        TAG_SPECIES=project.species_involved if project.organization else "MISSING!",
        TAG_LEAD_NAME=lead.get_full_name() if lead else "MISSING!",
        TAG_LEAD_NUMBER=lead.profile.phone if lead else "MISSING!",
        TAG_LEAD_EMAIL=lead.email if lead else "MISSING!",
        TAG_LEAD_POSITION=lead.profile.tposition if lead else "MISSING!",
        TAG_LEAD_CONTACT_INFO=contact_info if contact_info else "MISSING!",
        TAG_SECTION_HEAD_NAME=project.section.head.get_full_name() if project.section else "MISSING!",
        TAG_DIVISION_MANAGER_NAME=project.section.division.head.get_full_name() if project.section else "MISSING!",
        TAG_START_YEAR=project.years.first().start_date.strftime("%d/%m/%Y") if project.years.exists() else "MISSING!",
        TAG_END_YEAR=project.years.last().end_date.strftime("%d/%m/%Y") if project.years.exists() and project.years.last().end_date else "MISSING!",
        TAG_TEAM_DESCRIPTION=project.team_description,
        TAG_RATIONALE=project.rationale,
        TAG_OVERVIEW=project.overview,
        TAG_PRIORITIES=priorities,
        TAG_EXPERIMENTAL_PROTOCOL=project.experimental_protocol,
        TAG_DELIVERABLES=deliverables
    )

    for item in field_dict:
        # replace the tagged placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if item in paragraph.text:
                            try:
                                paragraph.text = paragraph.text.replace(item, field_dict[item])
                            except:
                                paragraph.text = "MISSING!"

        # replace the tagged placeholders in paragraphs
        for paragraph in document.paragraphs:
            if item in paragraph.text:
                try:
                    paragraph.text = paragraph.text.replace(item, field_dict[item])
                except:
                    paragraph.text = "MISSING!"

    # Find and populate the milestones and risk analysis table
    for table in document.tables:
        if "date/period" in table.rows[0].cells[0].paragraphs[0].text.lower():
            for milestone in models.Activity.objects.filter(type=1, project_year__project=project).order_by("target_date"):
                row = table.add_row()
                row.cells[0].paragraphs[0].text = milestone.target_date.strftime("%Y-%m-%d") if milestone.target_date else "MISSING!"
                row.cells[1].paragraphs[0].text = f"{milestone.name.upper()} - {milestone.description}"
                row.cells[2].paragraphs[0].text = milestone.responsible_party if milestone.responsible_party else "MISSING!"

        if "project risk analysis" in table.rows[0].cells[0].paragraphs[0].text.lower():
            for activity in models.Activity.objects.filter(project_year__project=project).order_by("target_date"):
                row = table.add_row()
                row.cells[0].paragraphs[0].text = f"{activity.name} ({activity.get_type_display()})"
                row.cells[1].paragraphs[0].text = activity.risk_description if activity.risk_description else "MISSING!"
                row.cells[2].paragraphs[0].text = f"{activity.likelihood}" if activity.likelihood else "MISSING!"
                row.cells[3].paragraphs[0].text = f"{activity.impact}" if activity.impact else "MISSING!"
                row.cells[4].paragraphs[0].text = activity.get_risk_rating_display() if activity.risk_rating else "MISSING!"
                row.cells[5].paragraphs[0].text = activity.mitigation_measures if activity.mitigation_measures else "MISSING!"
                row.cells[6].paragraphs[0].text = activity.responsible_party if activity.mitigation_measures else "MISSING!"

    document.save(target_file_path)

    return target_url


def generate_sar_workplan(year, region):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.xlsx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "sar_workplan_template.xlsx")

    year_txt = str(FiscalYear.objects.get(pk=year))

    # get all project years that are not in the following status: draft, not approved, cancelled
    # and that are a part of a project whose default funding source has an english name containing "csrf"
    qs = models.ProjectYear.objects.filter(project__default_funding_source__name__contains="SAR", fiscal_year=year)
    if region != "None":
        qs = qs.filter(project__section__division__branch__region_id=region)

    wb = load_workbook(filename=template_file_path)

    # to order workshees so the first sheet comes before the template sheet, rename the template and then copy the
    # renamed sheet, then rename the copy to template so it exists for other sheets to be created from
    ws = wb['template']
    ws.title = year_txt
    wb.copy_worksheet(ws).title = str("template")
    try:
        ws = wb[year_txt]
    except KeyError:
        print(year_txt, "is not a valid name of a worksheet")

    # start writing data at row 3 in the sheet
    row_count = 3
    for item in qs:

        ws['A' + str(row_count)].value = listrify([t.name for t in item.project.tags.all()])
        ws['J' + str(row_count)].value = ws['A' + str(row_count)].value

        ws['B' + str(row_count)].value = sum([c.amount for c in item.costs])
        ws['C' + str(row_count)].value = listrify(["{} {}".format(u.first_name, u.last_name) for u in item.get_project_leads_as_users()])

        ws['H' + str(row_count)].value = listrify([f.name for f in item.project.files.all()])
        ws['I' + str(row_count)].value = item.project.id

        if item.priorities:
            ws['K' + str(row_count)].value = html2text(item.priorities)

        ws['L' + str(row_count)].value = item.project.title

        if item.project.overview_html:
            ws['M' + str(row_count)].value = html2text(item.project.overview_html)

        activities = [html2text(act.description) for act in item.activities.filter(type=1)]
        ws['N' + str(row_count)].value = listrify(activities)

        activities = [html2text(act.description) for act in item.activities.filter(type=2)]
        ws['O' + str(row_count)].value = listrify(activities)

        row_count += 1

    wb.save(target_file_path)

    return target_url


def generate_acrdp_budget(project):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.xlsx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "acrdp_template.xlsx")

    om_category_2_cell = {
        # 1	Field Travel	Voyage sur le terrain(Travel)
        1: "24",
        # 2	DFO Business Travel (meeting etc.)	Voyage d'affaires du MPO (réunion, etc.) (Travel)
        2: "25",
        # 3	Training, domestic conferences	Formation, conférences domestiques (Travel)
        3: "26",
        # 4	Other Autre(Travel)
        4: "27",
        # 5	IM/IT - computers, hardware, software	GI / TI - ordinateurs, matériel informatique, logiciels (Equipment Purchase)
        5: "13",
        # 6	Lab Equipment	Équipement de laboratoire (Equipment Purchase)
        6: "14",
        # 7	Field Equipment	Équipement de terrain (Equipment Purchase)
        7: "15",
        # 8	Other	Autre (Equipment Purchase)
        8: "16",
        # 9	Office	Bureau (Material and Supplies)
        9: "21",
        # 10	Lab	Laboratoire (Material and Supplies)
        10: "19",
        # 11	Field	Terrain (Material and Supplies)
        11: "20",
        # 12	Other	Autre (Material and Supplies)
        12: "21",
        # 13	Students (FSWEP, Coop etc.)	Etudiants (PFETE, Coop, etc.) (Human Resources)
        13: "10",
        # 14	Post-Doctoral Candidates	Candidats postdoctoraux (Human Resources)
        14: "10",
        # 15	Contracts	Les contrats (Contracts, Leases, Services)
        15: "33",
        # 16	Translation	Traduction (Contracts, Leases, Services)
        16: "30",
        # 17	Publication costs	Frais de publication (Contracts, Leases, Services)
        17: "30",
        # 18	Vessels, Boats	Navires, Bateaux (Contracts, Leases, Services)
        18: "31",
        # 19	Facilities	Installations (Contracts, Leases, Services)
        19: "32",
        # 20	Other	Autre (Contracts, Leases, Services)
        20: "33",
        # 21	Fuel (e.g., boats)	Carburant (par exemple, bateaux) (Material and Supplies)
        21: "20",
        # 22	International travel for meetings, science collaboration and conferences	Voyages internationaux pour réunions, collaboration scientifique et conférences (Travel)
        22: "25",
        # 23	Equipment Maintenance	L'entretien de l'équipement (Other)
        23: "33",
    }

    capital_category_2_cell = {
        # 1, _("IM / IT - computers, hardware")),
        1: "13",
        # 2, _("Lab Equipment")),
        2: "14",
        # 3, _("Field Equipment")),
        3: "15",
        # 4, _("Other")),
        4: "16",
    }

    wb = load_workbook(filename=template_file_path)
    for year in project.years.all():
        try:
            ws = wb[str(year.fiscal_year)]
        except KeyError:
            print(str(year.fiscal_year), "is not a valid name of a worksheet")
        else:
            for cost in year.omcost_set.filter(funding_source__name__icontains="acrdp"):
                ref_cell = om_category_2_cell.get(cost.om_category_id)
                if ref_cell:
                    amount = ws['H' + ref_cell].value
                    description = ws['K' + ref_cell].value
                    ws['H' + ref_cell].value = nz(amount, 0) + nz(cost.amount, 0)
                    if not description:
                        ws['K' + ref_cell].value = cost.description
                    else:
                        ws['K' + ref_cell].value += "; " + cost.description

            for cost in year.capitalcost_set.filter(funding_source__name__icontains="acrdp"):
                ref_cell = capital_category_2_cell.get(cost.category)
                if ref_cell:
                    amount = ws['H' + ref_cell].value
                    description = ws['K' + ref_cell].value
                    ws['H' + ref_cell].value = nz(amount, 0) + nz(cost.amount, 0)
                    if not description:
                        ws['K' + ref_cell].value = cost.description
                    else:
                        ws['K' + ref_cell].value += "; " + cost.description

            for staff in year.staff_set.filter(funding_source__name__icontains="acrdp"):

                # determine the ref_cell
                if "student" in staff.employee_type.name.lower() or "post-doc" in staff.employee_type.name.lower() or staff.student_program:
                    ref_cell = 10
                elif not staff.level:
                    ref_cell = 7  # if we have no information about the level, we cannot choose a cell. Let's default to scientist
                elif "bi" in staff.level.name.lower():
                    ref_cell = 8  # if we have no information about the level, we cannot choose a cell. Let's default to scientist
                elif "eg" in staff.level.name.lower():
                    ref_cell = 9  # if we have no information about the level, we cannot choose a cell. Let's default to scientist
                elif "pc" in staff.level.name.lower() or "res" in staff.level.name.lower():
                    ref_cell = 7  # if we have no information about the level, we cannot choose a cell. Let's default to scientist
                else:
                    ref_cell = 7  # if we didn't catch it above, just through into the scientist bin again...
                ref_cell = str(ref_cell)
                # first need to determine if this is inkind or not:
                inkind = False
                if staff.amount == 0:  # in-kind
                    inkind = True

                amount = ws["H" + ref_cell].value
                description = ws["K" + ref_cell].value

                staff_description = staff.smart_name
                if staff.level: staff_description += f" ({staff.level})"
                if staff.duration_weeks: staff_description += f" @ {staff.duration_weeks} weeks"
                if inkind: staff_description += f" (in-kind)"

                ws["H" + ref_cell].value = nz(amount, 0) + nz(staff.amount, 0)
                if not description:
                    ws['K' + ref_cell].value = staff_description
                else:
                    ws['K' + ref_cell].value += "; " + staff_description

    wb.save(target_file_path)

    return target_url


def generate_csrf_submission_list(year, region):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.xlsx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "csrf_regional_submissions_template.xlsx")

    wb = load_workbook(filename=template_file_path)

    ws = wb["Submissions"]
    year_txt = str(FiscalYear.objects.get(pk=year))
    ws['A1'].value += year_txt

    # get all project years that are not in the following status: draft, not approved, cancelled
    # and that are a part of a project whose default funding source has an english name containing "csrf"
    qs = models.ProjectYear.objects.filter(project__default_funding_source__name__icontains="csrf", fiscal_year_id=year).filter(~Q(status__in=[1, 5, 9]))
    if region != "None":
        qs = qs.filter(project__section__division__branch__region_id=region)

    i = 3
    for item in qs:
        priority = None
        client_information = item.project.client_information
        if client_information:
            priority = client_information.csrf_priority

        theme = None
        pin = None
        if priority:
            theme = priority.csrf_sub_theme.csrf_theme.tname
            pin = priority.code

        leads = ""
        for l in item.get_project_leads_as_users():
            if l:
                leads += l.last_name + ", "
        if len(leads) > 0:
            leads = leads[:-2]

        # theme
        ws['A' + str(i)].value = nz(theme, " -----  ")
        # pin
        ws['B' + str(i)].value = nz(pin, " -----  ")
        # region
        ws['D' + str(i)].value = nz(item.project.section.division.branch.region.abbrev.upper(), " -----  ")
        # leads
        ws['E' + str(i)].value = nz(leads, " -----  ")
        # title
        ws['F' + str(i)].value = f'{item.project.title} ({item.project.id})'

        i += 1

    wb.save(target_file_path)

    return target_url


def generate_culture_committee_report():
    from publications import models as pi_models

    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_data_export_{}.xlsx".format(timezone.now().strftime("%Y-%m-%d"))
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)
    # create workbook and worksheets
    workbook = xlsxwriter.Workbook(target_file_path)

    # create formatting variables
    title_format = workbook.add_format({'bold': True, "align": 'normal', 'font_size': 24, })
    header_format = workbook.add_format(
        {'bold': True, 'border': 1, 'border_color': 'black', "align": 'normal', "text_wrap": True})
    total_format = workbook.add_format({'bold': True, "align": 'left', "text_wrap": True, 'num_format': '$#,##0'})
    normal_format = workbook.add_format({"align": 'left', "text_wrap": False, 'border': 1, 'border_color': 'black', })
    currency_format = workbook.add_format({'num_format': '#,##0.00'})
    date_format = workbook.add_format({'num_format': "yyyy-mm-dd", "align": 'left', })

    # get the dive list

    field_list = [
        "Project Id",
        "Title",
        "Description",
        "Years",
        "Keywords",
        "Leads",
        "Region",
        "Program / Funding Source",
        "Source",
    ]

    # get_cost_comparison_dict

    # define the header
    header = field_list
    title = "DM Apps Science Culture Committee Report"

    # define a worksheet
    my_ws = workbook.add_worksheet(name="projects")
    my_ws.write(0, 0, title, title_format)
    my_ws.write_row(2, 0, header, header_format)

    i = 3
    projects = models.Project.objects.filter(default_funding_source__is_competitive=True, years__status=4).distinct()
    for project in projects.order_by("id"):
        # create the col_max column to store the length of each header
        # should be a maximum column width to 100
        col_max = [len(str(d)) if len(str(d)) <= 100 else 100 for d in header]
        j = 0
        for field in field_list:
            my_val = None
            if "Project Id" in field:
                my_val = project.id
                my_ws.write(i, j, my_val, normal_format)

            elif "Title" in field:
                my_val = project.title
                my_ws.write(i, j, my_val, normal_format)

            elif "Description" in field:
                my_val = html2text(project.overview_html)
                my_ws.write(i, j, my_val, normal_format)

            elif "Years" in field:
                my_val = listrify([y.fiscal_year for y in project.years.filter(status=4)])
                my_ws.write(i, j, my_val, normal_format)

            elif "Keywords" in field:
                my_val = listrify([str(t) for t in project.tags.all()])
                my_ws.write(i, j, my_val, normal_format)

            elif "Leads" in field:
                my_val = listrify([str(staff) for staff in project.lead_staff.all()])
                my_ws.write(i, j, my_val, normal_format)

            elif "Region" in field:
                my_val = project.section.division.branch.region.tname
                my_ws.write(i, j, my_val, normal_format)

            elif "Program / Funding Source" in field:
                my_val = str(project.default_funding_source)
                my_ws.write(i, j, my_val, normal_format)

            elif field == "Source":
                my_val = "Project Planning"
                my_ws.write(i, j, my_val, normal_format)
            # adjust the width of the columns based on the max string length in each col
            ## replace col_max[j] if str length j is bigger than stored value

            # if new value > stored value... replace stored value
            if len(str(my_val)) > col_max[j]:
                if len(str(my_val)) < 50:
                    col_max[j] = len(str(my_val))
                else:
                    col_max[j] = 50
            j += 1
        i += 1

    archived_projects = pi_models.Project.objects.all()
    for project in archived_projects.order_by("id"):
        # create the col_max column to store the length of each header
        # should be a maximum column width to 100
        col_max = [len(str(d)) if len(str(d)) <= 100 else 100 for d in header]
        j = 0
        for field in field_list:
            my_val = None
            if "Project Id" in field:
                my_val = project.id
                my_ws.write(i, j, my_val, normal_format)

            elif "Title" in field:
                my_val = project.title
                my_ws.write(i, j, my_val, normal_format)

            elif "Description" in field:
                my_val = project.abstract
                my_ws.write(i, j, my_val, normal_format)

            elif "Years" in field:
                my_val = project.year
                my_ws.write(i, j, my_val, normal_format)

            elif "Keywords" in field:
                my_val = listrify([str(t) for t in project.theme.all()])
                my_ws.write(i, j, my_val, normal_format)

            elif "Leads" in field:
                my_val = listrify([str(staff) for staff in project.dfo_contact.all()])
                my_ws.write(i, j, my_val, normal_format)

            elif "Region" in field:
                if project.division.exists():
                    regions = Region.objects.filter(id__in=[p.branch.region.id for p in project.division.all()])
                    my_val = listrify([r.tname for r in regions])
                else:
                    my_val = "n/a"
                my_ws.write(i, j, my_val, normal_format)

            elif "Program / Funding Source" in field:
                my_val = listrify([str(item) for item in project.program_linkage.all()])
                my_ws.write(i, j, my_val, normal_format)

            elif field == "Source":
                my_val = "Project Inventory"
                my_ws.write(i, j, my_val, normal_format)

            # adjust the width of the columns based on the max string length in each col
            ## replace col_max[j] if str length j is bigger than stored value

            # if new value > stored value... replace stored value
            if len(str(my_val)) > col_max[j]:
                if len(str(my_val)) < 50:
                    col_max[j] = len(str(my_val))
                else:
                    col_max[j] = 50
            j += 1
        i += 1

        # set column widths
        for j in range(0, len(col_max)):
            my_ws.set_column(j, j, width=col_max[j] * 1.1)

    workbook.close()
    return target_url


def generate_project_status_summary(year, region):
    # Create the HttpResponse object with the appropriate CSV header.
    response = HttpResponse(content_type='text/csv')
    writer = csv.writer(response)
    status_choices = models.ProjectYear.status_choices

    headers = [
        'date',
        'fiscal_year',
        'region',
        'branch',
        'division',
        'section',
    ]
    for status in status_choices:
        headers.append(f'{slugify(status[1]).replace("-", "_")}_status_count')
    headers.append("total")

    qs = Section.objects.filter(projects2__years__fiscal_year_id=year).distinct()
    if region != "None":
        qs = qs.filter(division__branch__region_id=region)

    header_row = [header for header in headers]
    writer.writerow(header_row)

    for item in qs:
        data_row = [
            timezone.now(),
            str(FiscalYear.objects.get(pk=year)),
            item.division.branch.region.tname,
            item.division.branch.tname,
            item.division.tname,
            item.tname,
        ]
        for status in status_choices:
            data_row.append(models.ProjectYear.objects.filter(project__section=item, status=status[0]).count())
        data_row.append(models.ProjectYear.objects.filter(project__section=item).count())
        writer.writerow(data_row)

    return response


def generate_csrf_application(project, lang):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.docx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    if lang == "fr":
        template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "csrf_template_fr.docx")
    else:
        template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "csrf_template_en.docx")

    with open(template_file_path, 'rb') as f:
        source_stream = BytesIO(f.read())
    document = Document(source_stream)
    source_stream.close()

    priorities = str()
    for year in project.years.all():
        priorities += f'{year.fiscal_year}:\n{year.priorities}\n\n'

    references = str()
    for ref in project.references.all():
        references += f'{ref.short_citation}\n\n'

    project_length_1 = str()
    project_length_2 = str()
    year_count = project.years.count()
    project_length_str = f"{year_count} year{pluralize(year_count)}"
    if year_count <= 3:
        project_length_1 = project_length_str
    else:
        project_length_2 = project_length_str

    lead_name = str()
    lead_email = str()
    if project.lead_staff.exists():
        lead = project.lead_staff.first().user
        lead_name = lead.get_full_name()
        lead_email = lead.email

    field_work = _("No")
    for year in project.years.all():
        if year.has_field_component:
            field_work = _("Yes")
            break

    ship_time = _("No")
    for year in project.years.all():
        if year.coip_reference_id:
            ship_time = _("Yes") + f" (COIP ID: {year.coip_reference_id})"
            break

    risks = str()
    for activity in models.Activity.objects.filter(project_year__project=project):
        if activity.risk_description:
            risks += f'{activity.name.upper()}:\ni) {activity.risk_description}\nii) {activity.mitigation_measures}\n\n'

    data_management = str()
    for year in project.years.all():
        if year.has_data_component:
            data_management += f'{year.fiscal_year}:\n' \
                               f'i) {year.data_collected}\n' \
                               f'ii) {year.data_storage_plan}\n' \
                               f'iii) {year.data_products}\n\n'
    if len(data_management):
        data_management += "PLEASE REVIEW THIS SECTION!!!!"

    # COSTS
    years = project.years.order_by("fiscal_year")
    if year_count <= 3:
        years = project.years.all()[:3]

    cats = [
        ('salary', None),
        ('contracts', 5),
        ('equipment', 2),
        ('supplies', 3),
        ('travel', 1),
        ('vessel', None),  # 18
        ('other', 6),
        ('overhead', None),  # 24
        ('total', None),
    ]
    types = ["y1", 'y2', 'y3', 'total', 'detail']
    cost_dict = dict()
    for c in cats:
        for t in types:
            val = str() if t == "detail" else 0
            cost_dict[f'{c[0]}_{t}'] = val
    i = 1
    for year in years:
        # captial costs
        for cost in year.capitalcost_set.filter(amount__gt=0):
            cost_dict[f'equipment_y{i}'] += cost.amount
            cost_dict[f'equipment_total'] += cost.amount
            cost_description = f'{cost.project_year.fiscal_year} - {nz(cost.description, "MISSING DESCRIPTION")} = {currency(cost.amount, True)} (Captial)\n'
            cost_dict[f'equipment_detail'] += cost_description
            cost_dict[f'total_y{i}'] += cost.amount

        # rest of the costs
        for c in cats:
            cost_name = c[0]
            group = c[1]
            qs = None
            if not group:
                if cost_name == "salary":
                    # start with salary
                    for staff in year.staff_set.filter(amount__isnull=False, amount__gt=0):
                        cost_dict[f'salary_y{i}'] += staff.amount
                        cost_dict[f'salary_total'] += staff.amount
                        cost_dict[f'total_y{i}'] += staff.amount

                        staff_description = f'{year.fiscal_year} - {staff.smart_name}'
                        if staff.level: staff_description += f" ({staff.level})"
                        if staff.duration_weeks: staff_description += f" @ {staff.duration_weeks} weeks"
                        staff_description += f' = {currency(staff.amount, True)}\n'
                        cost_dict[f'salary_detail'] += staff_description

                elif cost_name == "overhead":
                    qs = year.omcost_set.filter(om_category__name__icontains="overhead", amount__gt=0)
                elif cost_name == "vessel":
                    qs = year.omcost_set.filter(om_category__name__icontains="vessels", amount__gt=0)

            else:
                qs = year.omcost_set.filter(om_category__group=group, amount__gt=0)
                if cost_name == "contracts":
                    qs = qs.filter(~Q(om_category__name__icontains="vessels"))
                elif cost_name == "other":
                    qs = qs.filter(~Q(om_category__name__icontains="overhead"))
            if qs:
                for cost in qs:
                    cost_dict[f'{cost_name}_y{i}'] += cost.amount
                    cost_dict[f'{cost_name}_total'] += cost.amount
                    cost_description = f'{cost.project_year.fiscal_year} - {nz(cost.description, "MISSING DESCRIPTION")} = {currency(cost.amount, True)} (O&M)\n'
                    cost_dict[f'{cost_name}_detail'] += cost_description
                    cost_dict[f'total_y{i}'] += cost.amount
        i += 1

    # calc the total_total
    cost_dict['total_total'] = cost_dict['total_y1'] + cost_dict['total_y2'] + cost_dict['total_y3']

    field_dict = dict(
        TAG_THEME=str(project.client_information.csrf_priority.csrf_sub_theme.csrf_theme) if project.client_information else "MISSING!",
        TAG_PRIORITY_CODE=str(project.client_information.csrf_priority.code) if project.client_information else "MISSING!",
        TAG_PRIORITY=str(project.client_information.csrf_priority) if project.client_information else "MISSING!",
        TAG_CLIENT_INFORMATION=str(project.client_information.description_en) if project.client_information else "MISSING!",
        TAG_SECOND_PRIORITY=str(project.second_priority),
        TAG_TITLE=project.title,
        TAG_PROJECT_LENGTH_1=project_length_1,
        TAG_PROJECT_LENGTH_2=project_length_2,
        TAG_LEAD_NAME=lead_name,
        TAG_LEAD_EMAIL=lead_email,
        TAG_REGION=str(project.section.division.branch.region.tname) if project.section else "MISSING!",
        TAG_OVERVIEW=project.overview,
        TAG_OBJECTIVES=project.objectives,
        TAG_PRIORITIES=priorities,
        TAG_INNOVATION=project.innovation,
        TAG_OTHER_FUNDING=project.other_funding,
        TAG_REFERENCES=references,
        TAG_FIELD_WORK=field_work,
        TAG_SHIP_TIME=ship_time,
        TAG_RISKS=risks,
        TAG_DATA_MANAGEMENT=data_management,

    )
    for c in cats:
        for t in types:
            val = str() if t == "detail" else 0
            field_dict[f'TAG_{c[0].upper()}_{t.upper()}'] = cost_dict[f'{c[0]}_{t}']

    for item in field_dict:
        # replace the tagged placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if item in paragraph.text:
                            try:
                                paragraph.text = paragraph.text.replace(item, str(field_dict[item]))
                            except Exception as E:
                                print(E, field_dict[item])
                                paragraph.text = "MISSING!"

        # replace the tagged placeholders in paragraphs
        for paragraph in document.paragraphs:
            if item in paragraph.text:
                try:
                    paragraph.text = paragraph.text.replace(item, field_dict[item])
                except:
                    paragraph.text = "MISSING!"

    # deal with staff
    for table in document.tables:
        if "theme" in table.rows[0].cells[0].paragraphs[0].text.lower():
            for staff in models.Staff.objects.filter(project_year__project=project).order_by("project_year"):
                row = table.add_row()
                row.cells[0].paragraphs[0].text = f'{staff.smart_name} ({staff.project_year.fiscal_year})'
                row.cells[1].paragraphs[
                    0].text = f'Role in the project: {staff.role} \nFTE Time: {int(nz(staff.duration_weeks, 0))} weeks ({round(int(nz(staff.duration_weeks, 0)) / 42 * 100)}%)\nKey Expertise: {staff.expertise}'
                row.cells[2].paragraphs[0].text = _("DFO-") + "[ADD REGION]"

            for c in models.Collaboration.objects.filter(project_year__project=project).order_by("project_year"):
                row = table.add_row()
                row.cells[0].paragraphs[0].text = f'{c.people} ({c.get_type_display()})'
                row.cells[1].paragraphs[0].text = c.notes
                row.cells[3].paragraphs[0].text = c.organization

        if "overview of the project" in table.rows[0].cells[0].paragraphs[0].text.lower():
            for activity in models.Activity.objects.filter(project_year__project=project).order_by("project_year"):
                row = table.add_row()
                row.cells[0].paragraphs[0].text = f'{activity.project_year.fiscal_year}'
                row.cells[1].paragraphs[
                    0].text = f'TYPE: {activity.get_type_display()} \nNAME: {activity.name}\nDESCRIPTION: {activity.description}\nRESPONSIBLE PARTIES: {activity.responsible_party}'

    document.save(target_file_path)

    return target_url


def generate_project_list(user, year, region, section):
    # Create the HttpResponse object with the appropriate CSV header.
    response = HttpResponse(content_type='text/csv')
    response.write(u'\ufeff'.encode('utf8'))  # BOM (optional...Excel needs it to open UTF-8 file properly)
    writer = csv.writer(response)
    status_choices = models.ProjectYear.status_choices

    fields = [
        'region',
        'division',
        'project.section|section',
        'project.id|Project Id',
        'fiscal_year',
        'project.title|title',
        'Overview',
        'Overview word count',
        'project.default_funding_source|Primary funding source',
        'project.functional_group|Functional group',
        'Project leads',
        'status',
        'updated_at|Last modified date',
        'modified_by|Last modified by',
        'Last modified description',
        'Activity count',
        'Staff count',
        'Sum of staff FTE (weeks)',
        'Sum of costs',
    ]

    if in_projects_admin_group(user):

        qs = ProjectYear.objects.filter(fiscal_year_id=year).distinct()
        if section != "None":
            qs = qs.filter(project__section_id=section)
        elif region != "None":
            qs = qs.filter(project__section__division__branch__region_id=region)
    else:
        sections = utils.get_manageable_sections(user)
        qs = ProjectYear.objects.filter(project__section__in=sections).distinct()

    header_row = [get_verbose_label(ProjectYear.objects.first(), header) for header in fields]
    writer.writerow(header_row)

    for obj in qs:
        data_row = list()
        for field in fields:
            if "division" in field:
                val = " ---"
                if obj.project.section:
                    val = obj.project.section.division.tname
            elif "region" in field:
                val = " ---"
                if obj.project.section:
                    val = obj.project.section.division.branch.region.tname
            elif "leads" in field:
                val = listrify(obj.get_project_leads_as_users())
            elif "updated_at" in field:
                val = obj.updated_at.strftime("%Y-%m-%d")
            elif "Last modified description" in field:
                val = naturaltime(obj.updated_at)
            elif field == "Overview":
                val = html2text(nz(obj.project.overview_html, ""))
            elif field == "Overview word count":
                val = len(html2text(nz(obj.project.overview_html, "")).split(" "))
            elif field == "Activity count":
                val = obj.activities.count()
            elif field == "Staff count":
                val = obj.staff_set.count()
            elif field == "Sum of staff FTE (weeks)":
                val = obj.staff_set.order_by("duration_weeks").aggregate(dsum=Sum("duration_weeks"))["dsum"]
            elif field == "Sum of costs":
                val = nz(obj.omcost_set.filter(amount__isnull=False).aggregate(dsum=Sum("amount"))["dsum"], 0) + \
                      nz(obj.capitalcost_set.filter(amount__isnull=False).aggregate(dsum=Sum("amount"))["dsum"], 0) + \
                      nz(obj.staff_set.filter(amount__isnull=False).aggregate(dsum=Sum("amount"))["dsum"], 0)
            else:
                val = get_field_value(obj, field)

            data_row.append(val)

        writer.writerow(data_row)

    return response


def generate_sara_application(project, lang):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.docx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    if lang == "fr":
        template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "sara_template.docx")
    else:
        template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "sara_template.docx")

    with open(template_file_path, 'rb') as f:
        source_stream = BytesIO(f.read())
    document = Document(source_stream)
    source_stream.close()

    year = project.years.last()
    priorities = str()
    priorities += f'{year.priorities}\n\n'

    milestones = str()
    for milestone in models.Activity.objects.filter(project_year=year, type=1):
        mystr = f'{date(milestone.target_date)} -> {milestone.name}: {milestone.description}. Responsible Parties: {milestone.responsible_party}'
        milestones += f'{mystr}\n'

    deliverables = str()
    for deliverable in models.Activity.objects.filter(project_year=year, type=2):
        mystr = f'{date(deliverable.target_date)} -> {deliverable.name}: {deliverable.description}. Responsible Parties: {deliverable.responsible_party}'
        deliverables += f'{mystr}\n'

    total_cost = 0
    om_costs = str()
    for cost in models.OMCost.objects.filter(project_year=year, funding_source__name__icontains="sara", amount__gt=0):
        mystr = f'{cost.om_category} --> {cost.description}  Amount: {currency(cost.amount)}'
        om_costs += f'{mystr}\n'
        total_cost += cost.amount

    capital_costs = str()
    for cost in models.CapitalCost.objects.filter(project_year=year, funding_source__name__icontains="sara", amount__gt=0):
        mystr = f'{cost.get_category_display()} --> {cost.description}  Amount: {currency(cost.amount)}'
        capital_costs += f'{mystr}\n'
        total_cost += cost.amount

    salary_costs = str()
    for staff in models.Staff.objects.filter(project_year=year, funding_source__name__icontains="sara", amount__gt=0):
        mystr = f'{nz(staff.smart_name)} ({nz(staff.level)}) --> Duration in weeks: {nz(staff.duration_weeks)}  Amount: {currency(staff.amount)}'
        salary_costs += f'{mystr}\n'
        total_cost += staff.amount

    field_dict = dict(
        TAG_TITLE=project.title,
        TAG_OM_COSTS=om_costs,
        TAG_CAPITAL_COSTS=capital_costs,
        TAG_SALARY_COSTS=salary_costs,
        TAG_TOTAL_COST=currency(total_cost, True),
        TAG_LEADS=listrify(project.lead_staff.all()),
        TAG_TAGS=listrify(project.tags.all()),
        TAG_OVERVIEW=project.overview,
        TAG_ADDITIONAL_NOTES=year.additional_notes,

        TAG_DELIVERABLES=deliverables,
        TAG_MILESTONES=milestones,
        TAG_PRIORITIES_METHODS=priorities,
        TAG_REPORTING=project.reporting_mechanism,
        TAG_FUNDING=project.future_funding_needs,
    )
    for item in field_dict:
        # replace the tagged placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if item in paragraph.text:
                            try:
                                paragraph.text = paragraph.text.replace(item, str(field_dict[item]))
                            except Exception as E:
                                print(E, field_dict[item])
                                paragraph.text = "MISSING!"

        # replace the tagged placeholders in paragraphs
        for paragraph in document.paragraphs:
            if item in paragraph.text:
                try:
                    paragraph.text = paragraph.text.replace(item, field_dict[item])
                except:
                    paragraph.text = "MISSING!"

    document.save(target_file_path)

    return target_url


def export_project_summary(request):
    # figure out the filename
    target_dir = os.path.join(settings.BASE_DIR, 'media', 'temp')
    target_file = "temp_export.xlsx"
    target_file_path = os.path.join(target_dir, target_file)
    target_url = os.path.join(settings.MEDIA_ROOT, 'temp', target_file)

    template_file_path = os.path.join(settings.BASE_DIR, 'projects2', 'static', "projects2", "projects_summary_template.xlsx")

    # get all project years using the same query as the management console
    qs = get_project_year_queryset(request)

    wb = load_workbook(filename=template_file_path)
    year_list = []
    # if the fiscal_year is set then we're only creating one worksheet for that year
    if 'fiscal_year' in request.GET and request.GET.get("fiscal_year"):
        year_list.append(request.GET.get("fiscal_year"))
    else:
        year_list = [y[0] for y in list(dict.fromkeys(qs.values_list('project__fiscal_years').distinct()))]

    for year in year_list:
        ws = wb['template']

        # to order worksheets so the first sheet comes before the template sheet, rename the template and then copy the
        # renamed sheet, then rename the copy to template so it exists for other sheets to be created from
        year_txt = str(shared_models.FiscalYear.objects.get(pk=year))
        ws.title = year_txt
        wb.copy_worksheet(ws).title = str("template")
        try:
            ws = wb[year_txt]
        except KeyError:
            print(year_txt, "is not a valid name of a worksheet")

        # start writing data at row 2 in the sheet
        row_count = 3
        for item in qs.filter(fiscal_year=year):
            project = item.project

            functional_group = project.functional_group if project.functional_group else None
            theme = functional_group.theme if functional_group and functional_group.theme else None
            primary_funding = project.default_funding_source if project.default_funding_source else None
            leads = ""
            for prj_lead in item.get_project_leads_as_users():
                if prj_lead:
                    if leads:
                        leads += ", "
                    leads += prj_lead.first_name + " " + prj_lead.last_name

            metadata = ""
            if project.updated_at:

                metadata += f"{naturaltime(project.updated_at)}"
                if project.modified_by:
                    metadata += f" by {project.modified_by}"

            priorites = item.priorities if item.priorities else None

            ws['A' + str(row_count)].value = project.pk
            ws['B' + str(row_count)].value = project.section.tname
            ws['C' + str(row_count)].value = project.title
            ws['D' + str(row_count)].value = html2text(nz(project.overview_html, ""))
            ws['E' + str(row_count)].value = nz(project.activity_type.tname, "") if project.activity_type else ""
            ws['F' + str(row_count)].value = nz(theme.tname, "") if theme else ""
            ws['G' + str(row_count)].value = nz(functional_group.tname, "") if functional_group else ""
            ws['H' + str(row_count)].value = str(primary_funding if primary_funding else "")
            ws['I' + str(row_count)].value = nz(str(project.start_date)[0:10], "")
            ws['J' + str(row_count)].value = nz(str(project.end_date)[0:10], "")
            ws['K' + str(row_count)].value = listrify(project.fiscal_years.all())
            ws['L' + str(row_count)].value = listrify(project.funding_sources.all())
            ws['M' + str(row_count)].value = nz(leads, "")
            ws['N' + str(row_count)].value = listrify(project.tags.all())
            ws['O' + str(row_count)].value = listrify(project.references.all(), "\n\r")
            ws['P' + str(row_count)].value = nz(metadata, "")
            # leave a column blank to separate general project info from the specif year project info.
            ws['R' + str(row_count)].value = html2text(nz(priorites, ""))
            ws['S' + str(row_count)].value = nz(item.requires_specialized_equipment, "")
            ws['T' + str(row_count)].value = nz(item.has_field_component, "")
            ws['U' + str(row_count)].value = nz(item.has_data_component, "")
            ws['V' + str(row_count)].value = nz(item.data_collected, "")
            ws['W' + str(row_count)].value = nz(item.data_products, "")
            ws['X' + str(row_count)].value = nz(item.data_management_needs, "")
            ws['Y' + str(row_count)].value = nz(item.has_lab_component, "")
            ws['Z' + str(row_count)].value = nz(item.it_needs, "")
            ws['AA' + str(row_count)].value = nz(item.additional_notes, "")
            ws['AB' + str(row_count)].value = nz(item.get_coding(), "")
            ws['AC' + str(row_count)].value = nz(item.submitted, "")
            ws['AD' + str(row_count)].value = nz(item.get_status_display(), "")
            # leave a column blank to separate specif year project info from O&M cost.
            ws['AF' + str(row_count)].value = nz(item.omcost_set.filter(funding_source__funding_source_type=1).aggregate(dsum=Sum("amount"))["dsum"], 0)
            ws['AG' + str(row_count)].value = nz(item.omcost_set.filter(funding_source__funding_source_type=2).aggregate(dsum=Sum("amount"))["dsum"], 0)
            ws['AH' + str(row_count)].value = nz(item.omcost_set.filter(funding_source__funding_source_type=3).aggregate(dsum=Sum("amount"))["dsum"], 0)
            ws['AI' + str(row_count)].value = nz(item.om_costs, "")
            row_count += 1

    wb.save(target_file_path)

    return target_url
